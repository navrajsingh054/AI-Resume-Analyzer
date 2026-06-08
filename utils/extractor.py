# utils/extractor.py
# ─────────────────────────────────────────────────────────────────────
# PURPOSE : Extract raw text from PDF and DOCX resume files.
#           Also detects section headings using font-size analysis.
# USED BY : app.py (main Streamlit UI)
# ─────────────────────────────────────────────────────────────────────

import fitz                      # PyMuPDF — reads PDF files
import docx                      # python-docx — reads Word files
import os                        # built-in — file path operations
import pandas as pd              # we'll store extracted data in a DataFrame


# ── CONSTANTS ────────────────────────────────────────────────────────
# Any text rendered at a font size >= this threshold is treated as a heading.
# Most resumes use 12-14pt for body text and 16-20pt for section titles.
HEADING_FONT_SIZE_THRESHOLD = 14.0


# ════════════════════════════════════════════════════════════════════
# SECTION 1 : PDF EXTRACTION
# ════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text from a PDF file as a single plain string.
    This is the simple version — just give me all the words.

    Args:
        file_path : full path to the .pdf file

    Returns:
        One big string of all text across all pages.
    """
    text = ""
    doc = fitz.open(file_path)        # load the entire PDF into memory

    for page_num in range(len(doc)):  # len(doc) = number of pages
        page = doc[page_num]          # get page object by index (0-based)
        text += page.get_text()       # get_text() returns plain string for this page
        text += "\n"                  # add blank line between pages so they don't merge

    doc.close()
    return text.strip()


def extract_structured_from_pdf(file_path: str) -> pd.DataFrame:
    """
    Advanced extraction: reads every span in the PDF and returns a
    DataFrame where each row is one line of text with metadata.

    This lets us detect headings (big font) vs body text (small font).

    Returns a DataFrame with columns:
        page    : page number (int)
        text    : the actual text of this line (str)
        font    : font name, e.g. "Arial-Bold" (str)
        size    : font size in points, e.g. 13.5 (float)
        is_bold : True if the span has bold flag set (bool)
        is_heading : True if size >= HEADING_FONT_SIZE_THRESHOLD (bool)
    """
    rows = []                         # we'll collect one dict per line here

    doc = fitz.open(file_path)

    for page_num in range(len(doc)):
        page = doc[page_num]

        # get_text("dict") returns the full block/line/span tree as a Python dict
        # instead of a plain string — this gives us font metadata too
        page_dict = page.get_text("dict")

        # page_dict["blocks"] is a list of block dicts
        # each block is either a text block (type=0) or image block (type=1)
        for block in page_dict["blocks"]:

            if block["type"] != 0:    # skip image blocks (type=1), we only want text
                continue

            # each block has "lines" — a list of line dicts
            for line in block["lines"]:

                # each line has "spans" — a list of span dicts
                # a span is one run of text with the same formatting
                line_text = ""
                line_size = 0.0
                line_font = ""
                line_bold = False

                for span in line["spans"]:
                    line_text += span["text"]   # concatenate all spans in the line

                    # take the largest font size seen in this line
                    if span["size"] > line_size:
                        line_size = span["size"]
                        line_font = span["font"]

                    # span["flags"] is a bitmask; bit 4 (value 16) = bold
                    # bitwise AND: if result != 0, bold flag is set
                    if span["flags"] & 16:
                        line_bold = True

                line_text = line_text.strip()
                if not line_text:             # skip completely empty lines
                    continue

                rows.append({
                    "page"       : page_num + 1,         # make it 1-based for humans
                    "text"       : line_text,
                    "font"       : line_font,
                    "size"       : round(line_size, 1),  # round to 1 decimal
                    "is_bold"    : line_bold,
                    "is_heading" : line_size >= HEADING_FONT_SIZE_THRESHOLD,
                })

    doc.close()

    # pd.DataFrame(rows) converts a list of dicts → a table
    # each key becomes a column name, each dict becomes a row
    df = pd.DataFrame(rows)
    return df


# ════════════════════════════════════════════════════════════════════
# SECTION 2 : DOCX EXTRACTION
# ════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a Word document as a single plain string.

    A .docx file is actually a ZIP archive of XML files.
    python-docx unpacks it and gives us paragraph objects.

    Args:
        file_path : full path to the .docx file

    Returns:
        One big string of all text.
    """
    doc = docx.Document(file_path)    # load the Word document

    # doc.paragraphs is a list of Paragraph objects
    # each paragraph has a .text property (the raw string)
    # and a .style.name property (e.g. "Heading 1", "Normal", "List Bullet")
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():         # skip blank paragraphs
            paragraphs.append(para.text)

    return "\n".join(paragraphs)


def extract_structured_from_docx(file_path: str) -> pd.DataFrame:
    """
    Advanced extraction for DOCX: uses paragraph styles to detect headings.

    In Word, headings are paragraphs with style names like:
        "Heading 1", "Heading 2", "Title"
    Body text has style names like:
        "Normal", "List Bullet", "List Paragraph"

    Returns a DataFrame with same columns as extract_structured_from_pdf()
    (except font/size, which Word doesn't expose easily at paragraph level).
    """
    doc = docx.Document(file_path)
    rows = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name  # e.g. "Heading 1" or "Normal"

        # check if this paragraph uses a heading style
        is_heading = style_name.startswith("Heading") or style_name == "Title"

        # para.runs is a list of Run objects (each run = same-formatted text)
        # we check if ANY run is bold
        is_bold = any(run.bold for run in para.runs if run.bold is not None)

        rows.append({
            "page"       : 1,          # Word docs don't have pages at paragraph level
            "text"       : text,
            "font"       : style_name, # use style name as the "font" for DOCX
            "size"       : 16.0 if is_heading else 12.0,  # estimate
            "is_bold"    : is_bold,
            "is_heading" : is_heading,
        })

    return pd.DataFrame(rows)


# ════════════════════════════════════════════════════════════════════
# SECTION 3 : SMART ROUTER — the only function other files need
# ════════════════════════════════════════════════════════════════════

def extract_text(file_path: str) -> str:
    """
    Detect file type and return plain text. Simple interface for other modules.
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported format '{ext}'. Upload PDF or DOCX only.")


def extract_structured(file_path: str) -> pd.DataFrame:
    """
    Detect file type and return structured DataFrame with metadata.
    Used by the section detector on Day 6.
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        return extract_structured_from_pdf(file_path)
    elif ext == ".docx":
        return extract_structured_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported format '{ext}'. Upload PDF or DOCX only.")


# ════════════════════════════════════════════════════════════════════
# SECTION 4 : SAVE UTILITY
# ════════════════════════════════════════════════════════════════════

def save_extracted_text(text: str, output_path: str) -> None:
    """
    Save extracted text to a .txt file so we can inspect it manually.
    Useful for debugging — open the .txt to check if extraction worked.

    Args:
        text        : the extracted text string
        output_path : where to save, e.g. "reports/extracted.txt"
    """
    # os.makedirs creates the folder if it doesn't exist
    # exist_ok=True means don't crash if the folder already exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    # open() in write mode ("w"), utf-8 handles special characters (accents, etc.)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"[extractor] Saved extracted text → {output_path}")